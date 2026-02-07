"""
Streamlit Web UI for Stock Research Assistant
Beautiful dashboard for Indian stock market analysis
"""
import html as html_module
import streamlit as st
import json
from datetime import datetime
import pandas as pd
import plotly.graph_objects as go
import yfinance as yf

# Must be first Streamlit command
st.set_page_config(
    page_title="Stock Research Assistant 🇮🇳",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Import tools after streamlit config
from tools.market_data import get_stock_price, get_stock_info, get_historical_data, get_index_data, get_trending_stocks
from tools.news_scraper import get_stock_news
from tools.analysis import calculate_technical_indicators, get_fundamental_metrics, analyze_price_action
from tools.institutional import get_fii_dii_data, get_bulk_block_deals
from config import NIFTY50_STOCKS, SECTORS


# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        background: linear-gradient(90deg, #FF9933, #FFFFFF, #138808);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        padding: 1rem;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
    }
    .positive { color: #00C851; font-weight: bold; }
    .negative { color: #ff4444; font-weight: bold; }
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        padding: 10px 20px;
        border-radius: 5px;
    }
    .report-box {
        background-color: rgba(255, 255, 255, 0.03);
        border: 1px solid rgba(255, 255, 255, 0.1);
        border-radius: 12px;
        padding: 2rem 2.5rem;
        margin: 1rem 0;
        line-height: 1.7;
    }
    .report-box h1 { font-size: 1.6rem; margin-top: 0; }
    .report-box h2 { font-size: 1.3rem; border-bottom: 1px solid rgba(255,255,255,0.1); padding-bottom: 0.4rem; }
    .report-box h3 { font-size: 1.1rem; }
    .report-box hr { border-color: rgba(255,255,255,0.08); }
</style>
""", unsafe_allow_html=True)


def format_number(num):
    """Format number in Indian style (lakhs, crores)."""
    if num is None or num == "N/A":
        return "N/A"
    try:
        num = float(num)
        if num >= 10_000_000:
            return f"₹{num/10_000_000:.2f} Cr"
        elif num >= 100_000:
            return f"₹{num/100_000:.2f} L"
        else:
            return f"₹{num:,.2f}"
    except (ValueError, TypeError):
        return str(num)


def get_trend_emoji(change):
    """Get emoji based on price change."""
    if change > 0:
        return "🟢"
    elif change < 0:
        return "🔴"
    return "⚪"


def _safe_val(value, prefix="", suffix=""):
    """Safely format a value for reports, handling None and N/A."""
    if value is None or value == "N/A":
        return "N/A"
    return f"{prefix}{value}{suffix}"


def _clean_report_markdown(report: str) -> str:
    """Clean up LLM-generated report for proper Streamlit rendering.

    Strips wrapping code fences that cause st.markdown to render
    the entire report as a monospace code block.
    """
    import re
    text = report.strip()
    # Remove wrapping ```markdown ... ``` or ``` ... ```
    text = re.sub(r'^```(?:markdown)?\s*\n', '', text)
    text = re.sub(r'\n```\s*$', '', text)
    return text.strip()


def generate_report_text(symbol: str, report_type: str, data: dict) -> str:
    """Generate a downloadable text report from analysis data."""
    if not data or not isinstance(data, dict):
        return f"Report for {symbol}: No data available."

    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S IST")

    if report_type == "technical":
        ma = data.get("moving_averages", {})
        momentum = data.get("momentum", {})
        volatility = data.get("volatility", {})
        sr = data.get("support_resistance", {})
        trend = data.get("trend", {})
        signals = data.get("signals", [])
        
        report = f"""
================================================================================
                    TECHNICAL ANALYSIS REPORT - {symbol}
================================================================================
Generated: {timestamp}
Source: Stock Research Assistant

OVERALL SIGNAL: {data.get('overall_signal', 'N/A')}
Signal Strength: {data.get('signal_strength', 'N/A')}

--------------------------------------------------------------------------------
MOVING AVERAGES
--------------------------------------------------------------------------------
SMA 20:  ₹{ma.get('sma_20', 'N/A')}
SMA 50:  ₹{ma.get('sma_50', 'N/A')}
SMA 200: ₹{ma.get('sma_200', 'N/A')}
Price vs SMA20: {ma.get('price_vs_sma20', 'N/A')}
Price vs SMA50: {ma.get('price_vs_sma50', 'N/A')}

--------------------------------------------------------------------------------
MOMENTUM INDICATORS
--------------------------------------------------------------------------------
RSI (14):        {momentum.get('rsi_14', 'N/A')} - {momentum.get('rsi_interpretation', 'N/A')}
MACD Line:       {momentum.get('macd_line', 'N/A')}
MACD Signal:     {momentum.get('macd_signal', 'N/A')}
Rate of Change:  {momentum.get('roc_10_day', 'N/A')}%

--------------------------------------------------------------------------------
VOLATILITY (BOLLINGER BANDS)
--------------------------------------------------------------------------------
Upper Band:  ₹{volatility.get('bollinger_upper', 'N/A')}
Middle Band: ₹{volatility.get('bollinger_middle', 'N/A')}
Lower Band:  ₹{volatility.get('bollinger_lower', 'N/A')}
BB Position: {volatility.get('bb_position', 'N/A')}
ATR %:       {volatility.get('atr_percent', 'N/A')}

--------------------------------------------------------------------------------
SUPPORT & RESISTANCE LEVELS
--------------------------------------------------------------------------------
Resistance 2: ₹{sr.get('resistance_2', 'N/A')}
Resistance 1: ₹{sr.get('resistance_1', 'N/A')}
Pivot:        ₹{sr.get('pivot', 'N/A')}
Support 1:    ₹{sr.get('support_1', 'N/A')}
Support 2:    ₹{sr.get('support_2', 'N/A')}

--------------------------------------------------------------------------------
TREND ANALYSIS
--------------------------------------------------------------------------------
Short-term:  {trend.get('short_term', 'N/A')}
Medium-term: {trend.get('medium_term', 'N/A')}
Long-term:   {trend.get('long_term', 'N/A')}
Golden Cross Active: {trend.get('golden_cross', 'N/A')}

--------------------------------------------------------------------------------
ACTIVE SIGNALS
--------------------------------------------------------------------------------
"""
        for sig in signals:
            report += f"- {sig.get('indicator', '')}: {sig.get('signal', '')} ({sig.get('strength', '')})\n"
        
        report += """
================================================================================
DISCLAIMER: This report is for educational purposes only. Not financial advice.
================================================================================
"""
        return report
    
    elif report_type == "fundamental":
        valuation = data.get("valuation") or {}
        profitability = data.get("profitability") or {}
        leverage = data.get("financial_health") or data.get("leverage") or {}
        growth = data.get("growth") or {}
        dividends = data.get("dividends") or {}
        size = data.get("size") or {}
        assessment = data.get("assessment") or []
        
        report = f"""
================================================================================
                   FUNDAMENTAL ANALYSIS REPORT - {symbol}
================================================================================
Generated: {timestamp}
Source: Stock Research Assistant

OVERALL RATING: {data.get('overall_rating', 'N/A')}
Valuation Status: {data.get('valuation_status', 'N/A')}

--------------------------------------------------------------------------------
VALUATION METRICS
--------------------------------------------------------------------------------
P/E Ratio:     {valuation.get('pe_ratio', 'N/A')}
P/B Ratio:     {valuation.get('pb_ratio', 'N/A')}
EV/EBITDA:     {valuation.get('ev_ebitda', 'N/A')}
Forward P/E:   {valuation.get('forward_pe', 'N/A')}
PEG Ratio:     {valuation.get('peg_ratio', 'N/A')}

--------------------------------------------------------------------------------
PROFITABILITY
--------------------------------------------------------------------------------
ROE:           {profitability.get('roe', 'N/A')}
ROA:           {profitability.get('roa', 'N/A')}
Gross Margin:  {profitability.get('gross_margin', 'N/A')}
Profit Margin: {profitability.get('profit_margin', 'N/A')}
Operating Margin: {profitability.get('operating_margin', 'N/A')}

--------------------------------------------------------------------------------
LEVERAGE & LIQUIDITY
--------------------------------------------------------------------------------
Debt/Equity:    {leverage.get('debt_to_equity', 'N/A')}
Current Ratio:  {leverage.get('current_ratio', 'N/A')}
Quick Ratio:    {leverage.get('quick_ratio', 'N/A')}

--------------------------------------------------------------------------------
GROWTH METRICS
--------------------------------------------------------------------------------
Earnings Growth:   {growth.get('earnings_growth', 'N/A')}
Revenue Growth:    {growth.get('revenue_growth', 'N/A')}
Quarterly EPS Growth: {growth.get('quarterly_earnings_growth', 'N/A')}

--------------------------------------------------------------------------------
DIVIDENDS
--------------------------------------------------------------------------------
Dividend Yield: {dividends.get('dividend_yield', 'N/A')}
Dividend Rate:  {dividends.get('dividend_rate', 'N/A')}
Payout Ratio:   {dividends.get('payout_ratio', 'N/A')}

--------------------------------------------------------------------------------
COMPANY SIZE
--------------------------------------------------------------------------------
Market Cap:       {size.get('market_cap', 'N/A')}
Enterprise Value: {size.get('enterprise_value', 'N/A')}
Revenue:          {size.get('revenue', 'N/A')}
EBITDA:           {size.get('ebitda', 'N/A')}

--------------------------------------------------------------------------------
KEY OBSERVATIONS
--------------------------------------------------------------------------------
"""
        for item in assessment:
            report += f"- {item.get('metric', '')}: {item.get('assessment', '')} [{item.get('impact', '')}]\n"
        
        report += """
================================================================================
DISCLAIMER: This report is for educational purposes only. Not financial advice.
================================================================================
"""
        return report
    
    return "Report not available"


def render_header():
    """Render the main header."""
    st.markdown('<h1 class="main-header">🇮🇳 Stock Research Assistant</h1>', unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: gray;'>AI-Powered Research for Indian Markets (NSE/BSE)</p>", unsafe_allow_html=True)
    st.divider()


def render_sidebar():
    """Render sidebar with stock selection."""
    with st.sidebar:
        st.image("https://upload.wikimedia.org/wikipedia/en/thumb/4/41/Flag_of_India.svg/1200px-Flag_of_India.svg.png", width=50)
        st.title("📊 Navigation")
        
        # Stock input
        st.subheader("🔍 Search Stock")
        symbol = st.text_input(
            "Enter Stock Symbol",
            placeholder="e.g., RELIANCE, TCS, INFY",
            help="Enter NSE stock symbol"
        ).upper().strip()
        
        # Trending stocks
        st.subheader("🔥 Trending Today")
        trending = get_trending_stocks()
        gainers = trending.get("gainers", [])
        losers = trending.get("losers", [])

        if gainers or losers:
            trend_tab1, trend_tab2 = st.tabs(["Top Gainers", "Top Losers"])

            with trend_tab1:
                for g in gainers[:5]:
                    sym = g.get("symbol", "")
                    chg = g.get("netPrice", 0)
                    st.markdown(f"**{sym}** :green[+{chg}%]")
                gainer_symbols = [g.get("symbol", "") for g in gainers[:5] if g.get("symbol")]
                selected_gainer = st.selectbox(
                    "Select gainer",
                    [""] + gainer_symbols,
                    format_func=lambda x: "Pick a stock..." if x == "" else x,
                    key="gainer_select",
                )
                if selected_gainer:
                    symbol = selected_gainer

            with trend_tab2:
                for ls in losers[:5]:
                    sym = ls.get("symbol", "")
                    chg = ls.get("netPrice", 0)
                    st.markdown(f"**{sym}** :red[{chg}%]")
                loser_symbols = [ls.get("symbol", "") for ls in losers[:5] if ls.get("symbol")]
                selected_loser = st.selectbox(
                    "Select loser",
                    [""] + loser_symbols,
                    format_func=lambda x: "Pick a stock..." if x == "" else x,
                    key="loser_select",
                )
                if selected_loser:
                    symbol = selected_loser
        else:
            # Fallback to hardcoded list
            st.subheader("⚡ Quick Select")
            selected_stock = st.selectbox(
                "Popular Stocks",
                [""] + NIFTY50_STOCKS[:20],
                format_func=lambda x: "Select a stock..." if x == "" else x,
            )
            if selected_stock:
                symbol = selected_stock
        
        # Sector filter
        st.subheader("🏭 Browse by Sector")
        selected_sector = st.selectbox(
            "Sector",
            [""] + list(SECTORS.keys()),
            format_func=lambda x: "All Sectors" if x == "" else x
        )
        
        if selected_sector:
            sector_stocks = SECTORS.get(selected_sector, [])
            sector_stock = st.selectbox("Stocks", [""] + sector_stocks)
            if sector_stock:
                symbol = sector_stock
        
        st.divider()
        
        # Market status
        now = datetime.now()
        hour = now.hour
        if now.weekday() >= 5:
            status = "🔴 Market Closed (Weekend)"
        elif hour < 9 or (hour == 9 and now.minute < 15):
            status = "🟡 Pre-Market"
        elif hour < 15 or (hour == 15 and now.minute <= 30):
            status = "🟢 Market Open"
        else:
            status = "🔴 Market Closed"
        
        st.markdown(f"**Market Status:** {status}")
        st.markdown(f"**Last Updated:** {now.strftime('%H:%M:%S IST')}")
        
        st.divider()
        st.caption("⚠️ For educational purposes only. Not financial advice.")
        
        return symbol


def render_market_overview():
    """Render market overview section."""
    st.subheader("🏦 Market Overview")
    
    try:
        indices_data = json.loads(get_index_data.run("ALL"))
        
        col1, col2, col3, col4 = st.columns(4)
        
        index_cols = [
            ("NIFTY50", "NIFTY 50", col1),
            ("SENSEX", "SENSEX", col2),
            ("BANKNIFTY", "BANK NIFTY", col3),
            ("NIFTYIT", "NIFTY IT", col4),
        ]
        
        for key, name, col in index_cols:
            if key in indices_data:
                data = indices_data[key]
                value = data.get("value", 0)
                change = data.get("change", 0)
                change_pct = data.get("change_percent", 0)
                
                with col:
                    st.metric(
                        label=name,
                        value=f"{value:,.2f}",
                        delta=f"{change:+,.2f} ({change_pct:+.2f}%)"
                    )
    except Exception as e:
        st.warning(f"Could not fetch market data: {e}")


def _fetch_chart_data(symbol: str, period: str) -> pd.DataFrame:
    """Fetch OHLCV data from yfinance for charting.

    Returns a DataFrame with a **timezone-naive IST** DatetimeIndex.
    Intraday timestamps are shifted to interval-end so the last candle
    of the day reads 3:30 PM (market close).
    For 1D/1W the last candle's Close is patched with the official daily
    closing price so the chart endpoint matches the header price.
    """
    from tools.market_data import _get_nse_symbol
    from datetime import timedelta, time as dt_time

    MARKET_CLOSE = dt_time(15, 30)

    # (yfinance period, interval, timedelta to shift candle to end-of-interval)
    period_map = {
        "1D": ("5d", "5m", timedelta(minutes=5)),
        "1W": ("5d", "5m", timedelta(minutes=5)),
        "1M": ("1mo", "30m", timedelta(minutes=30)),
        "3M": ("3mo", "1h", timedelta(hours=1)),
        "6M": ("6mo", "1d", None),
        "1Y": ("1y", "1d", None),
        "5Y": ("5y", "1d", None),
    }
    yf_period, interval, shift = period_map.get(period, ("1y", "1d", None))

    try:
        ticker = yf.Ticker(_get_nse_symbol(symbol))
        df = ticker.history(period=yf_period, interval=interval)

        if df.empty:
            return df

        # Convert to naive IST (yfinance already returns Asia/Kolkata)
        if df.index.tz is not None:
            df.index = df.index.tz_convert("Asia/Kolkata").tz_localize(None)

        # Shift intraday candles to interval-end timestamps,
        # capping at 15:30 (market close) so the last bar is correct.
        if shift is not None:
            new_idx = df.index + shift
            capped = []
            for ts in new_idx:
                if ts.time() > MARKET_CLOSE:
                    ts = ts.replace(hour=15, minute=30, second=0, microsecond=0)
                capped.append(ts)
            df.index = pd.DatetimeIndex(capped)

        # For 1D, keep only the most recent trading day
        if period == "1D":
            last_date = df.index[-1].date()
            df = df[df.index.date == last_date]

        # Patch the last candle of each intraday day with the official
        # daily close so the chart endpoint matches the header price.
        # (15-min candles miss the closing auction.)
        if shift is not None:
            try:
                daily = ticker.history(period=yf_period, interval="1d")
                if daily.index.tz is not None:
                    daily.index = daily.index.tz_convert("Asia/Kolkata").tz_localize(None)
                daily_close_map = {d.date(): row["Close"] for d, row in daily.iterrows()}
                for i in range(len(df) - 1, -1, -1):
                    ts = df.index[i]
                    if ts.time() == MARKET_CLOSE:
                        official = daily_close_map.get(ts.date())
                        if official is not None:
                            df.iloc[i, df.columns.get_loc("Close")] = official
            except Exception:
                pass  # best-effort; fall back to candle data

        return df
    except Exception:
        return pd.DataFrame()


def _render_range_bar(label: str, low: float, high: float, current: float):
    """Render a horizontal range bar with a triangle marker via HTML/CSS."""
    if high <= low or high == 0:
        return

    pct = max(0, min(100, ((current - low) / (high - low)) * 100))

    html = f"""
    <div style="margin: 0.6rem 0;">
      <div style="display:flex; justify-content:space-between; font-size:0.82rem; color:#aaa; margin-bottom:2px;">
        <span>{label} Low — ₹{low:,.2f}</span>
        <span>₹{high:,.2f} — {label} High</span>
      </div>
      <div style="position:relative; height:6px; background:linear-gradient(90deg,#ff4444 0%,#ffcc00 50%,#00C851 100%); border-radius:3px;">
        <div style="position:absolute; top:-6px; left:{pct}%; transform:translateX(-50%);">
          <div style="width:0; height:0; border-left:6px solid transparent; border-right:6px solid transparent; border-top:8px solid white;"></div>
        </div>
      </div>
    </div>
    """
    st.markdown(html, unsafe_allow_html=True)


def render_stock_overview(symbol: str):
    """Render Groww-style stock overview with price header, range bars, fundamentals grid, and chart."""
    try:
        with st.spinner(f"Fetching data for {symbol}..."):
            price_data = json.loads(get_stock_price.run(symbol))
            info_data = json.loads(get_stock_info.run(symbol))

        if "error" in price_data:
            st.error(f"Stock not found: {symbol}")
            return None

        # ── Section 1: Company header ──
        change = price_data.get("change", 0)
        change_pct = price_data.get("change_percent", 0)
        color = "#00C851" if change >= 0 else "#ff4444"
        sign = "+" if change >= 0 else ""
        company_name = info_data.get("company_name", symbol)

        st.markdown(
            f"""
            <div style="margin-bottom:0.5rem;">
              <span style="font-size:1.1rem; color:#aaa;">{html_module.escape(str(company_name))}</span>
              <span style="font-size:0.85rem; color:#888; margin-left:8px;">{html_module.escape(str(info_data.get('sector', '')))} · {html_module.escape(str(info_data.get('industry', '')))}</span>
            </div>
            <div style="display:flex; align-items:baseline; gap:12px; flex-wrap:wrap;">
              <span style="font-size:2.4rem; font-weight:700;">₹{price_data.get('current_price', 0):,.2f}</span>
              <span style="font-size:1.1rem; color:{color}; font-weight:600;">{sign}{change:,.2f} ({sign}{change_pct:.2f}%)</span>
            </div>
            """,
            unsafe_allow_html=True,
        )

        # ── Section 2: Range bars ──
        day_low = price_data.get("low", 0)
        day_high = price_data.get("high", 0)
        w52_low = price_data.get("52_week_low", 0)
        w52_high = price_data.get("52_week_high", 0)
        current = price_data.get("current_price", 0)

        if isinstance(day_low, (int, float)) and isinstance(day_high, (int, float)) and day_high > day_low:
            _render_range_bar("Today's", day_low, day_high, current)
        if isinstance(w52_low, (int, float)) and isinstance(w52_high, (int, float)) and w52_high > w52_low:
            _render_range_bar("52W", w52_low, w52_high, current)

        st.markdown("<div style='margin-top:1rem;'></div>", unsafe_allow_html=True)

        # ── Section 3: Fundamentals grid ──
        st.markdown("#### Fundamentals")

        def _fmt(val, prefix="", suffix=""):
            if val is None or val == "N/A":
                return "—"
            return f"{prefix}{val}{suffix}"

        fund_items = [
            ("Market Cap", format_number(info_data.get("market_cap"))),
            ("ROE", _fmt(info_data.get("roe"), suffix="%")),
            ("P/E Ratio", _fmt(info_data.get("pe_ratio"))),
            ("EPS", _fmt(info_data.get("eps"), prefix="₹")),
            ("P/B Ratio", _fmt(info_data.get("pb_ratio"))),
            ("Div. Yield", _fmt(info_data.get("dividend_yield"), suffix="%")),
            ("Debt/Equity", _fmt(info_data.get("debt_to_equity"))),
            ("Book Value", _fmt(info_data.get("book_value"), prefix="₹")),
        ]

        # 2-column grid, 4 rows
        for row_start in range(0, len(fund_items), 2):
            cols = st.columns(2)
            for idx, col in enumerate(cols):
                item_idx = row_start + idx
                if item_idx < len(fund_items):
                    label, value = fund_items[item_idx]
                    col.markdown(
                        f"<div style='padding:6px 0; border-bottom:1px solid rgba(255,255,255,0.06);'>"
                        f"<span style='color:#aaa; font-size:0.85rem;'>{label}</span><br>"
                        f"<span style='font-size:1.05rem; font-weight:500;'>{value}</span></div>",
                        unsafe_allow_html=True,
                    )

        # ── Section 4: Interactive price chart ──
        st.markdown("<div style='margin-top:1.5rem;'></div>", unsafe_allow_html=True)

        periods = ["1D", "1W", "1M", "3M", "6M", "1Y", "5Y"]
        selected_period = st.radio(
            "Chart Period",
            periods,
            index=5,  # default 1Y
            horizontal=True,
            key=f"chart_period_{symbol}",
        )

        with st.spinner("Loading chart..."):
            chart_df = _fetch_chart_data(symbol, selected_period)

        if not chart_df.empty:
            line_color = "#00C851" if change >= 0 else "#ff4444"
            fill_color = "rgba(0,200,81,0.10)" if change >= 0 else "rgba(255,68,68,0.10)"
            is_intraday = selected_period in ("1D", "1W", "1M")

            # Hover: show 12-hour time for intraday, date for daily
            if is_intraday:
                hover_tpl = "₹%{y:,.2f}<extra>%{x|%d %b %Y, %I:%M %p} IST</extra>"
            else:
                hover_tpl = "₹%{y:,.2f}<extra>%{x|%d %b %Y}</extra>"

            fig = go.Figure()

            # Invisible baseline trace at the data min so fill doesn't go to y=0
            y_floor = float(chart_df["Close"].min())
            fig.add_trace(
                go.Scatter(
                    x=chart_df.index,
                    y=[y_floor] * len(chart_df),
                    mode="lines",
                    line=dict(width=0),
                    showlegend=False,
                    hoverinfo="skip",
                )
            )

            fig.add_trace(
                go.Scatter(
                    x=chart_df.index,
                    y=chart_df["Close"],
                    mode="lines",
                    line=dict(color=line_color, width=2),
                    fill="tonexty",
                    fillcolor=fill_color,
                    showlegend=False,
                    hovertemplate=hover_tpl,
                )
            )

            # Pad y-axis 5% above and below the data range
            y_min = float(chart_df["Close"].min())
            y_max = float(chart_df["Close"].max())
            y_pad = (y_max - y_min) * 0.05 or 1

            xaxis_opts: dict = dict(showgrid=False)
            if is_intraday:
                xaxis_opts["tickformat"] = "%-I:%M %p"

            fig.update_layout(
                height=380,
                margin=dict(l=0, r=0, t=10, b=0),
                xaxis=xaxis_opts,
                yaxis=dict(
                    showgrid=True,
                    gridcolor="rgba(255,255,255,0.05)",
                    side="right",
                    range=[y_min - y_pad, y_max + y_pad],
                    tickprefix="₹",
                ),
                plot_bgcolor="rgba(0,0,0,0)",
                paper_bgcolor="rgba(0,0,0,0)",
                hovermode="x unified",
            )
            st.plotly_chart(fig, use_container_width=True, key=f"chart_{symbol}_{selected_period}")
        else:
            st.info("Chart data not available for this period.")

        return price_data, info_data

    except Exception as e:
        st.error(f"Error fetching stock data: {e}")
        return None


def render_technical_analysis(symbol: str):
    """Render technical analysis tab."""
    try:
        with st.spinner("Calculating technical indicators..."):
            tech_data = json.loads(calculate_technical_indicators.run(symbol))
        
        if "error" in tech_data:
            st.error(tech_data["error"])
            return
        
        # Overall signal
        signal = tech_data.get("overall_signal", "NEUTRAL")
        signal_colors = {"BULLISH": "green", "BEARISH": "red", "NEUTRAL": "orange"}
        signal_emoji = {"BULLISH": "🟢", "BEARISH": "🔴", "NEUTRAL": "🟡"}
        
        st.markdown(f"### {signal_emoji.get(signal, '⚪')} Overall Signal: **:{signal_colors.get(signal, 'gray')}[{signal}]**")
        st.caption(f"Signal Strength: {tech_data.get('signal_strength', 'N/A')}")
        
        st.divider()
        
        # Indicators in columns
        col1, col2, col3 = st.columns(3)
        
        ma = tech_data.get("moving_averages", {})
        momentum = tech_data.get("momentum", {})
        volatility = tech_data.get("volatility", {})
        
        with col1:
            st.markdown("#### 📈 Moving Averages")
            st.markdown(f"""
            | Indicator | Value |
            |-----------|-------|
            | SMA 20 | ₹{ma.get('sma_20', 'N/A')} |
            | SMA 50 | ₹{ma.get('sma_50', 'N/A')} |
            | SMA 200 | ₹{ma.get('sma_200', 'N/A')} |
            | vs SMA20 | {ma.get('price_vs_sma20', 'N/A')} |
            | vs SMA50 | {ma.get('price_vs_sma50', 'N/A')} |
            """)
        
        with col2:
            st.markdown("#### 📊 Momentum")
            rsi = momentum.get('rsi_14', 0)
            rsi_color = "red" if rsi > 70 else ("green" if rsi < 30 else "gray")
            st.markdown(f"""
            | Indicator | Value |
            |-----------|-------|
            | RSI (14) | :{rsi_color}[{rsi}] |
            | Interpretation | {momentum.get('rsi_interpretation', 'N/A')} |
            | MACD Line | {momentum.get('macd_line', 'N/A')} |
            | MACD Signal | {momentum.get('macd_signal', 'N/A')} |
            | ROC (10d) | {momentum.get('roc_10_day', 'N/A')}% |
            """)
        
        with col3:
            st.markdown("#### 📉 Volatility")
            st.markdown(f"""
            | Indicator | Value |
            |-----------|-------|
            | BB Upper | ₹{volatility.get('bollinger_upper', 'N/A')} |
            | BB Middle | ₹{volatility.get('bollinger_middle', 'N/A')} |
            | BB Lower | ₹{volatility.get('bollinger_lower', 'N/A')} |
            | BB Position | {volatility.get('bb_position', 'N/A')} |
            | ATR % | {volatility.get('atr_percent', 'N/A')} |
            """)
        
        st.divider()
        
        # Support/Resistance and Trend
        col1, col2 = st.columns(2)
        
        sr = tech_data.get("support_resistance", {})
        trend = tech_data.get("trend", {})
        
        with col1:
            st.markdown("#### 🎯 Support & Resistance")
            st.markdown(f"""
            | Level | Price |
            |-------|-------|
            | Resistance 2 | ₹{sr.get('resistance_2', 'N/A')} |
            | Resistance 1 | ₹{sr.get('resistance_1', 'N/A')} |
            | **Pivot** | **₹{sr.get('pivot', 'N/A')}** |
            | Support 1 | ₹{sr.get('support_1', 'N/A')} |
            | Support 2 | ₹{sr.get('support_2', 'N/A')} |
            """)
        
        with col2:
            st.markdown("#### 📈 Trend Analysis")
            
            def trend_badge(t):
                if t == "Bullish":
                    return "🟢 Bullish"
                elif t == "Bearish":
                    return "🔴 Bearish"
                return "⚪ N/A"
            
            st.markdown(f"""
            | Timeframe | Trend |
            |-----------|-------|
            | Short-term | {trend_badge(trend.get('short_term', 'N/A'))} |
            | Medium-term | {trend_badge(trend.get('medium_term', 'N/A'))} |
            | Long-term | {trend_badge(trend.get('long_term', 'N/A'))} |
            """)
            
            gc = trend.get('golden_cross')
            if gc is True:
                st.success("✨ Golden Cross Active (Bullish)")
            elif gc is False:
                st.warning("💀 Death Cross Active (Bearish)")
        
        # Trading Signals
        st.divider()
        st.markdown("#### ⚡ Active Signals")
        
        signals = tech_data.get("signals", [])
        if signals:
            for sig in signals:
                indicator = sig.get("indicator", "")
                signal_text = sig.get("signal", "")
                strength = sig.get("strength", "")
                
                if "Buy" in signal_text or "Bullish" in signal_text:
                    st.success(f"**{indicator}:** {signal_text} ({strength})")
                elif "Sell" in signal_text or "Bearish" in signal_text:
                    st.error(f"**{indicator}:** {signal_text} ({strength})")
                else:
                    st.info(f"**{indicator}:** {signal_text} ({strength})")
        else:
            st.info("No strong signals at the moment")
        
        # Download button for technical report
        st.divider()
        report_text = generate_report_text(symbol, "technical", tech_data)
        col1, col2 = st.columns([1, 3])
        with col1:
            st.download_button(
                label="📥 Download Technical Report",
                data=report_text,
                file_name=f"{symbol}_technical_analysis_{datetime.now().strftime('%Y-%m-%d')}.txt",
                mime="text/plain",
                use_container_width=True
            )
    
    except Exception as e:
        st.error(f"Error in technical analysis: {e}")


def render_fundamental_analysis(symbol: str):
    """Render fundamental analysis tab."""
    try:
        with st.spinner("Fetching fundamental metrics..."):
            fund_data = json.loads(get_fundamental_metrics.run(symbol))
        
        if "error" in fund_data:
            st.error(fund_data["error"])
            return
        
        # Rating header
        rating = fund_data.get("overall_rating", "N/A")
        rating_colors = {
            "STRONG BUY": "green",
            "BUY": "green", 
            "HOLD": "orange",
            "SELL": "red",
            "STRONG SELL": "red",
        }
        rating_emojis = {
            "STRONG BUY": "🟢🟢",
            "BUY": "🟢",
            "HOLD": "🟡",
            "SELL": "🔴",
            "STRONG SELL": "🔴🔴",
        }
        
        st.markdown(f"### {rating_emojis.get(rating, '⚪')} Fundamental Rating: **:{rating_colors.get(rating, 'gray')}[{rating}]**")
        st.caption(f"Score: {fund_data.get('score', 'N/A')} | Rating: {fund_data.get('rating_percentage', 'N/A')}")
        
        st.divider()
        
        # Metrics in columns
        col1, col2, col3 = st.columns(3)
        
        val = fund_data.get("valuation", {})
        prof = fund_data.get("profitability", {})
        health = fund_data.get("financial_health", {})
        
        with col1:
            st.markdown("#### 💰 Valuation")
            st.markdown(f"""
            | Metric | Value |
            |--------|-------|
            | P/E Ratio | {val.get('pe_ratio', 'N/A')} |
            | Forward P/E | {val.get('forward_pe', 'N/A')} |
            | P/B Ratio | {val.get('pb_ratio', 'N/A')} |
            | P/S Ratio | {val.get('ps_ratio', 'N/A')} |
            | EV/EBITDA | {val.get('ev_ebitda', 'N/A')} |
            | PEG Ratio | {val.get('peg_ratio', 'N/A')} |
            """)
        
        with col2:
            st.markdown("#### 📈 Profitability")
            st.markdown(f"""
            | Metric | Value |
            |--------|-------|
            | ROE | {prof.get('roe', 'N/A')} |
            | ROA | {prof.get('roa', 'N/A')} |
            | Profit Margin | {prof.get('profit_margin', 'N/A')} |
            | Operating Margin | {prof.get('operating_margin', 'N/A')} |
            | Gross Margin | {prof.get('gross_margin', 'N/A')} |
            """)
        
        with col3:
            st.markdown("#### 🏦 Financial Health")
            debt_status = health.get('debt_status', 'N/A')
            debt_color = "green" if debt_status == "Low" else ("orange" if debt_status == "Moderate" else "red")
            st.markdown(f"""
            | Metric | Value |
            |--------|-------|
            | Debt/Equity | {health.get('debt_to_equity', 'N/A')} |
            | Current Ratio | {health.get('current_ratio', 'N/A')} |
            | Quick Ratio | {health.get('quick_ratio', 'N/A')} |
            | Debt Status | :{debt_color}[{debt_status}] |
            """)
        
        st.divider()
        
        # Growth and Dividends
        col1, col2, col3 = st.columns(3)
        
        growth = fund_data.get("growth", {})
        div = fund_data.get("dividends", {})
        size = fund_data.get("size", {})
        
        with col1:
            st.markdown("#### 📊 Growth")
            st.markdown(f"""
            | Metric | Value |
            |--------|-------|
            | Earnings Growth | {growth.get('earnings_growth', 'N/A')} |
            | Revenue Growth | {growth.get('revenue_growth', 'N/A')} |
            | Quarterly EPS Growth | {growth.get('quarterly_earnings_growth', 'N/A')} |
            """)
        
        with col2:
            st.markdown("#### 💵 Dividends")
            st.markdown(f"""
            | Metric | Value |
            |--------|-------|
            | Dividend Yield | {div.get('dividend_yield', 'N/A')} |
            | Dividend Rate | {div.get('dividend_rate', 'N/A')} |
            | Payout Ratio | {div.get('payout_ratio', 'N/A')} |
            """)
        
        with col3:
            st.markdown("#### 📏 Company Size")
            st.markdown(f"""
            | Metric | Value |
            |--------|-------|
            | Market Cap | {size.get('market_cap', 'N/A')} |
            | Enterprise Value | {size.get('enterprise_value', 'N/A')} |
            | Revenue | {size.get('revenue', 'N/A')} |
            | EBITDA | {size.get('ebitda', 'N/A')} |
            """)
        
        # Assessment
        assessment = fund_data.get("assessment", [])
        if assessment:
            st.divider()
            st.markdown("#### 📋 Key Observations")
            for item in assessment:
                metric = item.get("metric", "")
                assess = item.get("assessment", "")
                impact = item.get("impact", "")
                
                if impact == "Positive":
                    st.success(f"✅ **{metric}:** {assess}")
                elif impact == "Negative":
                    st.error(f"❌ **{metric}:** {assess}")
                else:
                    st.info(f"ℹ️ **{metric}:** {assess}")
        
        # Download button for fundamental report
        st.divider()
        report_text = generate_report_text(symbol, "fundamental", fund_data)
        col1, col2 = st.columns([1, 3])
        with col1:
            st.download_button(
                label="📥 Download Fundamental Report",
                data=report_text,
                file_name=f"{symbol}_fundamental_analysis_{datetime.now().strftime('%Y-%m-%d')}.txt",
                mime="text/plain",
                use_container_width=True
            )
    
    except Exception as e:
        st.error(f"Error in fundamental analysis: {e}")


def render_news(symbol: str):
    """Render news tab."""
    try:
        with st.spinner("Fetching latest news..."):
            news_data = json.loads(get_stock_news.run(symbol, 5))
        
        articles = news_data.get("articles", [])
        sources_status = news_data.get("sources_status", {})
        
        # Source status
        st.markdown("#### 📰 News Sources")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            if sources_status.get("et_rss") == "success":
                st.success("✅ ET RSS")
            else:
                st.warning("⚠️ ET RSS")

        with col2:
            if sources_status.get("economic_times") == "success":
                st.success("✅ Economic Times")
            else:
                st.warning("⚠️ Economic Times")

        with col3:
            if sources_status.get("google_news") == "success":
                st.success("✅ Google News")
            else:
                st.warning("⚠️ Google News")
        
        st.divider()
        
        if not articles:
            st.info(f"No recent news found for {symbol}")
            return
        
        st.markdown(f"#### 📰 Latest News ({len(articles)} articles)")
        
        for i, article in enumerate(articles, 1):
            title = article.get("title", "No title")
            source = article.get("source", "Unknown")
            url = article.get("url", "")
            published = article.get("published", "")
            summary = article.get("summary", "")
            
            with st.expander(f"**{i}. {title}**", expanded=i<=3):
                st.caption(f"Source: {source} | {published}")
                if summary:
                    st.write(summary)
                if url:
                    st.markdown(f"[Read full article →]({url})")
    
    except Exception as e:
        st.error(f"Error fetching news: {e}")


def render_institutional(symbol: str):
    """Render institutional activity tab."""
    try:
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### 🏦 FII/DII Activity")
            with st.spinner("Fetching FII/DII data..."):
                fii_data = json.loads(get_fii_dii_data.run())
            
            if "error" not in fii_data:
                fii = fii_data.get("fii", {})
                dii = fii_data.get("dii", {})
                combined = fii_data.get("combined", {})
                
                if fii.get("buy_value_cr"):
                    st.markdown(f"""
                    **FII Activity:**
                    - Buy: ₹{fii.get('buy_value_cr', 0):,.0f} Cr
                    - Sell: ₹{fii.get('sell_value_cr', 0):,.0f} Cr
                    - Net: ₹{fii.get('net_value_cr', 0):,.0f} Cr ({fii.get('activity', 'N/A')})
                    
                    **DII Activity:**
                    - Buy: ₹{dii.get('buy_value_cr', 0):,.0f} Cr
                    - Sell: ₹{dii.get('sell_value_cr', 0):,.0f} Cr
                    - Net: ₹{dii.get('net_value_cr', 0):,.0f} Cr ({dii.get('activity', 'N/A')})
                    
                    **Market Sentiment:** {combined.get('market_sentiment', 'N/A')}
                    """)
                else:
                    st.info("Check NSE/BSE websites for latest FII/DII data")
            else:
                st.warning("Could not fetch FII/DII data")
        
        with col2:
            st.markdown(f"#### 📊 Bulk/Block Deals for {symbol}")
            with st.spinner("Fetching deals..."):
                deals_data = json.loads(get_bulk_block_deals.run(symbol))
            
            deals = deals_data.get("deals", [])
            
            has_real_deals = (
                deals
                and isinstance(deals[0], dict)
                and "note" not in deals[0]
            )
            if has_real_deals:
                for deal in deals[:5]:
                    st.markdown(f"""
                    - **{deal.get('stock', 'N/A')}**: {deal.get('deal_type', 'N/A')}
                      - Client: {deal.get('client', 'N/A')}
                      - Qty: {deal.get('quantity', 'N/A')}
                      - Price: {deal.get('price', 'N/A')}
                    """)
            else:
                st.info(f"No recent bulk/block deals for {symbol}")
                st.markdown("""
                **What to look for:**
                - Large investor activity (FIIs, DIIs, HNIs)
                - Promoter transactions
                - Private equity exits
                """)
    
    except Exception as e:
        st.error(f"Error: {e}")


def _generate_word_report(symbol: str, markdown_report: str, report_time: str) -> bytes:
    """Convert a markdown report into a Word (.docx) document and return bytes."""
    import io
    import re
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Clean code fences from LLM output
    cleaned = _clean_report_markdown(markdown_report)

    doc = Document()

    # Professional styling - set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading(f"Stock Research Report — {symbol}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Timestamp
    ts_para = doc.add_paragraph()
    ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts_run = ts_para.add_run(f"Generated: {report_time}")
    ts_run.font.size = Pt(10)
    ts_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph("")  # spacer

    # Table accumulation state machine
    table_buffer = []

    def flush_table():
        if table_buffer:
            _add_word_table(doc, list(table_buffer))
            table_buffer.clear()

    # Parse markdown line-by-line
    for line in cleaned.split("\n"):
        stripped = line.strip()

        # Skip residual ``` lines
        if stripped.startswith("```"):
            continue

        # Detect pipe-table lines
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            continue
        else:
            flush_table()

        if not stripped:
            doc.add_paragraph("")
            continue

        # Headings - strip inline markdown from heading text
        if stripped.startswith("### "):
            doc.add_heading(_strip_inline_md(stripped[4:]), level=3)
        elif stripped.startswith("## "):
            doc.add_heading(_strip_inline_md(stripped[3:]), level=2)
        elif stripped.startswith("# "):
            doc.add_heading(_strip_inline_md(stripped[2:]), level=1)
        elif stripped.startswith("---"):
            _add_horizontal_rule(doc)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            para = doc.add_paragraph(style="List Bullet")
            _add_formatted_runs(para, text)
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            para = doc.add_paragraph(style="List Number")
            _add_formatted_runs(para, text)
        else:
            para = doc.add_paragraph()
            _add_formatted_runs(para, stripped)

    # Flush any remaining table lines
    flush_table()

    # Disclaimer footer
    doc.add_paragraph("")
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "DISCLAIMER: This report is for educational purposes only. Not financial advice."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)
    run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _sanitize_for_pdf(text: str) -> str:
    """Replace non-latin1 characters with ASCII equivalents for fpdf."""
    replacements = {
        '\u20b9': 'Rs.',   # ₹
        '\u2014': '--',    # —
        '\u2013': '-',     # –
        '\u2018': "'",     # '
        '\u2019': "'",     # '
        '\u201c': '"',     # "
        '\u201d': '"',     # "
        '\u2022': '-',     # •
        '\u2026': '...',   # …
        '\u00d7': 'x',     # ×
    }
    for char, repl in replacements.items():
        text = text.replace(char, repl)
    # Fallback: drop any remaining non-latin1 characters
    return text.encode('latin-1', errors='replace').decode('latin-1')


def _generate_pdf_report(symbol: str, markdown_report: str, report_time: str) -> bytes:
    """Convert a markdown report into a PDF and return bytes."""
    import re
    from fpdf import FPDF

    cleaned = _sanitize_for_pdf(_clean_report_markdown(markdown_report))

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Title
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(27, 42, 74)  # dark navy
    pdf.cell(w=0, text=f"Stock Research Report -- {symbol}",
             new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(4)

    # Timestamp
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(w=0, text=f"Generated: {report_time}",
             new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    # Table accumulation
    table_buffer: list[str] = []

    def _pdf_row_height(col_widths, cells, line_h=5):
        """Calculate the tallest cell height for a row with text wrapping."""
        max_lines = 1
        for i, text in enumerate(cells):
            w = col_widths[i] - 2  # padding
            text_w = pdf.get_string_width(text)
            n_lines = max(1, int(text_w / w) + 1) if w > 0 else 1
            if n_lines > max_lines:
                max_lines = n_lines
        return max_lines * line_h

    def _pdf_draw_row(col_widths, cells, row_h, fill_color=None, is_header=False):
        """Draw a single table row with multi_cell wrapping."""
        x_start = pdf.get_x()
        y_start = pdf.get_y()

        # Check if row fits on page, otherwise add new page
        if y_start + row_h > pdf.h - pdf.b_margin:
            pdf.add_page()
            y_start = pdf.get_y()

        if fill_color:
            pdf.set_fill_color(*fill_color)

        for i, text in enumerate(cells):
            pdf.set_xy(x_start + sum(col_widths[:i]), y_start)
            # Draw cell border + fill as a rect, then overlay text
            pdf.rect(x_start + sum(col_widths[:i]), y_start,
                     col_widths[i], row_h, style="DF" if fill_color else "D")
            pdf.set_xy(x_start + sum(col_widths[:i]) + 1, y_start + 1)
            pdf.multi_cell(w=col_widths[i] - 2, h=5,
                           text=text, border=0, fill=False,
                           align="C" if is_header else "L",
                           new_x="RIGHT", new_y="TOP")

        pdf.set_xy(x_start, y_start + row_h)

    def flush_pdf_table():
        if not table_buffer:
            return
        rows = []
        for line in table_buffer:
            stripped = line.strip().replace('|', '\t').replace('-', '').replace(':', '').strip()
            if not stripped:
                continue  # separator
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            rows.append(cells)

        if not rows:
            table_buffer.clear()
            return

        num_cols = len(rows[0])
        page_width = pdf.w - pdf.l_margin - pdf.r_margin
        col_widths = [page_width / num_cols] * num_cols

        # Header row
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(255, 255, 255)
        header_cells = [_strip_inline_md(c) for c in rows[0]]
        h_height = _pdf_row_height(col_widths, header_cells)
        _pdf_draw_row(col_widths, header_cells, h_height,
                      fill_color=(27, 42, 74), is_header=True)

        # Data rows
        pdf.set_font("Helvetica", "", 9)
        for r_idx, row in enumerate(rows[1:]):
            data_cells = [_strip_inline_md(c) for c in row[:num_cols]]
            r_height = _pdf_row_height(col_widths, data_cells)
            fill = (242, 242, 242) if r_idx % 2 == 1 else (255, 255, 255)
            pdf.set_text_color(0, 0, 0)
            _pdf_draw_row(col_widths, data_cells, r_height, fill_color=fill)

        pdf.ln(4)
        table_buffer.clear()

    for line in cleaned.split("\n"):
        stripped = line.strip()

        if stripped.startswith("```"):
            continue

        # Table lines
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(stripped)
            continue
        else:
            flush_pdf_table()

        if not stripped:
            pdf.ln(3)
            continue

        clean_text = _strip_inline_md(stripped)

        if stripped.startswith("### "):
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(27, 42, 74)
            pdf.cell(w=0, text=clean_text[4:] if clean_text.startswith("### ") else _strip_inline_md(stripped[4:]),
                     new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
        elif stripped.startswith("## "):
            pdf.set_font("Helvetica", "B", 14)
            pdf.set_text_color(27, 42, 74)
            pdf.cell(w=0, text=_strip_inline_md(stripped[3:]),
                     new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)
        elif stripped.startswith("# "):
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(27, 42, 74)
            pdf.cell(w=0, text=_strip_inline_md(stripped[2:]),
                     new_x="LMARGIN", new_y="NEXT")
            pdf.ln(4)
        elif stripped.startswith("---"):
            y = pdf.get_y()
            pdf.set_draw_color(170, 170, 170)
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(4)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(0, 0, 0)
            bullet_text = _strip_inline_md(stripped[2:])
            # Color BUY/SELL keywords
            if "BUY" in bullet_text.upper():
                pdf.set_text_color(0, 128, 0)
            elif "SELL" in bullet_text.upper():
                pdf.set_text_color(200, 0, 0)
            pdf.cell(w=6, h=6, text="-")
            pdf.multi_cell(w=0, h=6, text=f" {bullet_text}",
                           new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)
        elif re.match(r"^\d+\.\s", stripped):
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(0, 0, 0)
            num_match = re.match(r"^(\d+\.)\s(.*)", stripped)
            if num_match:
                num_prefix = num_match.group(1)
                body = _strip_inline_md(num_match.group(2))
                pdf.cell(w=8, h=6, text=num_prefix)
                pdf.multi_cell(w=0, h=6, text=f" {body}",
                               new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(0, 0, 0)
            text_out = _strip_inline_md(stripped)
            if "BUY" in text_out.upper():
                pdf.set_text_color(0, 128, 0)
            elif "SELL" in text_out.upper():
                pdf.set_text_color(200, 0, 0)
            pdf.multi_cell(w=0, h=6, text=text_out,
                           new_x="LMARGIN", new_y="NEXT")
            pdf.set_text_color(0, 0, 0)

    flush_pdf_table()

    # Disclaimer
    pdf.ln(8)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(w=0, text="DISCLAIMER: This report is for educational purposes only. Not financial advice.",
             new_x="LMARGIN", new_y="NEXT", align="C")

    return bytes(pdf.output())


def _strip_inline_md(text: str) -> str:
    """Strip inline markdown markers (bold/italic) for use in headings."""
    import re
    # Remove bold-italic (***), bold (**), italic (*)
    text = re.sub(r'\*{3}(.+?)\*{3}', r'\1', text)
    text = re.sub(r'\*{2}(.+?)\*{2}', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    return text


def _add_formatted_runs(paragraph, text: str, base_font_size=None):
    """Add text with inline markdown (bold/italic) and BUY/SELL coloring.

    Parses ***bold-italic***, **bold**, *italic* into proper Word runs.
    Also colors BUY/SELL keywords green/red.
    """
    import re
    from docx.shared import RGBColor

    # First split on BUY/SELL keywords
    keyword_parts = re.split(r"(STRONG BUY|STRONG SELL|BUY|SELL)", text)

    for kpart in keyword_parts:
        upper = kpart.strip().upper()
        if upper in ("BUY", "STRONG BUY"):
            run = paragraph.add_run(kpart)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 128, 0)
            if base_font_size:
                run.font.size = base_font_size
            continue
        if upper in ("SELL", "STRONG SELL"):
            run = paragraph.add_run(kpart)
            run.bold = True
            run.font.color.rgb = RGBColor(200, 0, 0)
            if base_font_size:
                run.font.size = base_font_size
            continue

        # Parse inline markdown: ***bold-italic***, **bold**, *italic*
        tokens = re.split(
            r'(\*{3}.+?\*{3}|\*{2}.+?\*{2}|\*[^*]+?\*)', kpart
        )
        for token in tokens:
            if not token:
                continue
            if token.startswith('***') and token.endswith('***'):
                run = paragraph.add_run(token[3:-3])
                run.bold = True
                run.italic = True
            elif token.startswith('**') and token.endswith('**'):
                run = paragraph.add_run(token[2:-2])
                run.bold = True
            elif token.startswith('*') and token.endswith('*') and len(token) > 2:
                run = paragraph.add_run(token[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(token)
            if base_font_size:
                run.font.size = base_font_size


def _add_colored_run(paragraph, text: str):
    """Add text to a paragraph, coloring BUY/SELL keywords."""
    _add_formatted_runs(paragraph, text)


def _set_cell_shading(cell, color_hex: str):
    """Set background shading on a Word table cell via XML."""
    from docx.oxml.ns import qn
    from lxml import etree

    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')


def _add_word_table(doc, table_lines: list):
    """Convert markdown pipe-table lines into a proper Word table.

    Expects lines like: | Header1 | Header2 | ... |
    Second line is separator (|---|---|) and is skipped.
    """
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if len(table_lines) < 2:
        return

    def parse_row(line):
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        return cells

    header_cells = parse_row(table_lines[0])
    num_cols = len(header_cells)

    # Collect data rows, skipping separator line (|---|---|)
    data_rows = []
    for line in table_lines[1:]:
        stripped = line.strip().replace('|', '').replace('-', '').replace(':', '').strip()
        if not stripped:
            continue  # separator line
        data_rows.append(parse_row(line))

    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, text in enumerate(header_cells):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(_strip_inline_md(text))
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(cell, '1B2A4A')

    # Data rows with alternating shading
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, text in enumerate(row_data[:num_cols]):
            cell = table.rows[1 + row_idx].cells[col_idx]
            cell.text = ''
            _add_formatted_runs(cell.paragraphs[0], _strip_inline_md(text), base_font_size=Pt(10))
            if row_idx % 2 == 1:
                _set_cell_shading(cell, 'F2F2F2')

    doc.add_paragraph('')  # spacer after table


def _add_horizontal_rule(doc):
    """Add a horizontal rule as a paragraph with a bottom border."""
    from docx.oxml.ns import qn
    from lxml import etree

    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn('w:pBdr'))
    bottom = etree.SubElement(pBdr, qn('w:bottom'))
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')


def render_ai_analysis(symbol: str):
    """Render AI analysis tab."""
    st.markdown("#### 🤖 AI-Powered Full Analysis")
    
    st.info("""
    **Full AI Analysis** uses 6 specialized AI agents to create a comprehensive research report:
    
    1. 📊 **Market Data Analyst** - Collects all market data
    2. 📰 **News Analyst** - Analyzes news sentiment
    3. 💰 **Fundamental Analyst** - Deep financial analysis
    4. 📈 **Technical Analyst** - Chart patterns & indicators
    5. 🎯 **Investment Strategist** - Synthesizes recommendations
    6. 📝 **Report Writer** - Creates the final report
    
    **Requirements:**
    - Mistral AI API key in your `.env` file
    - Takes 2-3 minutes to complete
    """)
    
    if st.button(f"🚀 Run Full AI Analysis for {symbol}", type="primary", use_container_width=True):
        try:
            from config import settings
            if not settings.mistral_api_key:
                st.error("❌ Mistral API key not configured. Please add MISTRAL_API_KEY to your .env file.")
                return
            
            from crews.research_crew import analyze_stock_sync
            
            with st.spinner(f"🔬 AI agents analyzing {symbol}... This may take 2-3 minutes."):
                report = analyze_stock_sync(symbol, "full")
            
            st.success("✅ Analysis complete!")

            # Store report in session state for download
            st.session_state[f"report_{symbol}"] = report
            st.session_state[f"report_time_{symbol}"] = datetime.now().strftime("%Y-%m-%d_%H-%M")

            st.divider()
            cleaned = _clean_report_markdown(report)
            with st.container(border=True):
                st.markdown(cleaned)
            
        except Exception as e:
            st.error(f"Error during AI analysis: {e}")
    
    # Show stored report and download buttons
    if f"report_{symbol}" in st.session_state:
        report = st.session_state[f"report_{symbol}"]
        report_time = st.session_state.get(f"report_time_{symbol}", datetime.now().strftime("%Y-%m-%d_%H-%M"))

        with st.expander(f"📄 View Report ({report_time})", expanded=True):
            cleaned = _clean_report_markdown(report)
            st.markdown(cleaned)

        st.markdown("#### 📥 Download Report")

        col1, col2 = st.columns(2)

        with col1:
            pdf_bytes = _generate_pdf_report(symbol, report, report_time)
            st.download_button(
                label="📄 Download as PDF",
                data=pdf_bytes,
                file_name=f"{symbol}_research_report_{report_time}.pdf",
                mime="application/pdf",
                use_container_width=True,
            )

        with col2:
            word_bytes = _generate_word_report(symbol, report, report_time)
            st.download_button(
                label="📝 Download as Word",
                data=word_bytes,
                file_name=f"{symbol}_research_report_{report_time}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )


def main():
    """Main application."""
    render_header()
    
    # Sidebar and get selected symbol
    symbol = render_sidebar()
    
    # Market overview at top
    render_market_overview()
    
    st.divider()
    
    # Main content
    if symbol:
        result = render_stock_overview(symbol)
        
        if result:
            st.divider()
            
            # Tabs for different analyses
            tab1, tab2, tab3, tab4, tab5 = st.tabs([
                "📈 Technical",
                "💰 Fundamentals", 
                "📰 News",
                "🏦 Institutional",
                "🤖 AI Analysis"
            ])
            
            with tab1:
                render_technical_analysis(symbol)
            
            with tab2:
                render_fundamental_analysis(symbol)
            
            with tab3:
                render_news(symbol)
            
            with tab4:
                render_institutional(symbol)
            
            with tab5:
                render_ai_analysis(symbol)
    else:
        # Landing page when no stock selected
        st.markdown("""
        ### 👋 Welcome to Stock Research Assistant!
        
        **Get started:**
        1. Enter a stock symbol in the sidebar (e.g., `RELIANCE`, `TCS`, `INFY`)
        2. Or select from popular stocks dropdown
        3. Explore Technical, Fundamental, and News analysis
        
        **Features:**
        - 📊 Real-time price data from NSE/BSE
        - 📈 Technical indicators (RSI, MACD, Bollinger Bands)
        - 💰 Fundamental metrics (P/E, ROE, Debt analysis)
        - 📰 News from Economic Times, Google News
        - 🏦 FII/DII activity tracking
        - 🤖 Full AI-powered research reports
        """)
        
        # Show some popular stocks
        st.subheader("🔥 Popular Stocks")
        
        cols = st.columns(5)
        popular = ["RELIANCE", "TCS", "HDFCBANK", "INFY", "ICICIBANK"]
        
        for i, stock in enumerate(popular):
            with cols[i]:
                try:
                    price_data = json.loads(get_stock_price.run(stock))
                    change = price_data.get("change_percent", 0)
                    trend = "🟢" if change >= 0 else "🔴"
                    st.metric(
                        f"{trend} {stock}",
                        f"₹{price_data.get('current_price', 0):,.2f}",
                        f"{change:+.2f}%"
                    )
                except Exception:
                    st.metric(stock, "—", help="Price data temporarily unavailable")


if __name__ == "__main__":  # pragma: no cover
    main()
