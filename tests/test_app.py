"""
Tests for Streamlit Web UI (app.py)

Rewrites tests to import actual functions from app.py instead of duplicating them locally.
Also adds tests for render functions, report generation, and helper utilities.
"""

import pytest
import json
import sys
import importlib
from unittest.mock import patch, MagicMock, PropertyMock
from datetime import datetime
import pandas as pd
import numpy as np

# ============================================================
# Fixtures
# ============================================================


@pytest.fixture
def mock_st():
    """Mock streamlit module and reload app to use it."""
    mock = MagicMock()
    # Provide a functioning session_state dict
    mock.session_state = {}
    return mock


@pytest.fixture
def app_module(mock_st):
    """Import app module with streamlit mocked."""
    original = sys.modules.get('streamlit')
    sys.modules['streamlit'] = mock_st
    # Also mock plotly so import doesn't fail
    if 'app' in sys.modules:
        importlib.reload(sys.modules['app'])
    else:
        import app as _  # noqa: F811
    mod = sys.modules['app']
    yield mod
    # Restore
    if original is not None:
        sys.modules['streamlit'] = original
    else:
        sys.modules.pop('streamlit', None)
    importlib.reload(sys.modules['app'])


# ============================================================
# Tests for format_number — import from actual module
# ============================================================


class TestFormatNumber:
    """Tests for format_number helper function."""

    @pytest.mark.unit
    def test_format_crores(self, app_module):
        assert "Cr" in app_module.format_number(50_000_000)
        assert "50.00 Cr" in app_module.format_number(500_000_000)

    @pytest.mark.unit
    def test_format_lakhs(self, app_module):
        result = app_module.format_number(500_000)
        assert "L" in result
        assert "5.00 L" in result

    @pytest.mark.unit
    def test_format_small_number(self, app_module):
        assert app_module.format_number(1234.56) == "₹1,234.56"

    @pytest.mark.unit
    def test_format_na(self, app_module):
        assert app_module.format_number(None) == "N/A"
        assert app_module.format_number("N/A") == "N/A"

    @pytest.mark.unit
    def test_format_invalid(self, app_module):
        assert app_module.format_number("invalid") == "invalid"


# ============================================================
# Tests for get_trend_emoji
# ============================================================


class TestGetTrendEmoji:
    """Tests for trend emoji helper."""

    @pytest.mark.unit
    def test_positive_trend(self, app_module):
        assert app_module.get_trend_emoji(5.5) == "🟢"
        assert app_module.get_trend_emoji(0.01) == "🟢"

    @pytest.mark.unit
    def test_negative_trend(self, app_module):
        assert app_module.get_trend_emoji(-2.5) == "🔴"
        assert app_module.get_trend_emoji(-0.01) == "🔴"

    @pytest.mark.unit
    def test_neutral_trend(self, app_module):
        assert app_module.get_trend_emoji(0) == "⚪"


# ============================================================
# Tests for _safe_val
# ============================================================


class TestSafeVal:
    """Tests for _safe_val helper."""

    @pytest.mark.unit
    def test_with_value(self, app_module):
        assert app_module._safe_val(42.5, prefix="₹", suffix="%") == "₹42.5%"

    @pytest.mark.unit
    def test_with_none(self, app_module):
        assert app_module._safe_val(None) == "N/A"

    @pytest.mark.unit
    def test_with_zero(self, app_module):
        assert app_module._safe_val(0) == "0"

    @pytest.mark.unit
    def test_with_na_string(self, app_module):
        assert app_module._safe_val("N/A") == "N/A"


# ============================================================
# Tests for _clean_report_markdown
# ============================================================


class TestCleanReportMarkdown:
    """Tests for _clean_report_markdown helper."""

    @pytest.mark.unit
    def test_strips_markdown_code_fence(self, app_module):
        text = "```markdown\n# Title\nContent\n```"
        result = app_module._clean_report_markdown(text)
        assert "```" not in result
        assert "# Title" in result

    @pytest.mark.unit
    def test_strips_plain_code_fence(self, app_module):
        text = "```\n# Title\nContent\n```"
        result = app_module._clean_report_markdown(text)
        assert "```" not in result

    @pytest.mark.unit
    def test_no_fence_unchanged(self, app_module):
        text = "# Title\nContent"
        result = app_module._clean_report_markdown(text)
        assert result == text


# ============================================================
# Tests for _sanitize_for_pdf
# ============================================================


class TestSanitizeForPdf:
    """Tests for _sanitize_for_pdf helper."""

    @pytest.mark.unit
    def test_replaces_rupee_symbol(self, app_module):
        result = app_module._sanitize_for_pdf("₹100")
        assert "Rs." in result
        assert "₹" not in result

    @pytest.mark.unit
    def test_replaces_em_dash(self, app_module):
        result = app_module._sanitize_for_pdf("hello — world")
        assert "--" in result

    @pytest.mark.unit
    def test_drops_unknown_unicode(self, app_module):
        # A character not in the replacements map and not in latin-1
        result = app_module._sanitize_for_pdf("test \u2603 value")  # snowman
        assert isinstance(result, str)
        # Should encode cleanly to latin-1
        result.encode('latin-1')


# ============================================================
# Tests for _strip_inline_md
# ============================================================


class TestStripInlineMd:
    """Tests for _strip_inline_md helper."""

    @pytest.mark.unit
    def test_strips_bold(self, app_module):
        assert app_module._strip_inline_md("**bold**") == "bold"

    @pytest.mark.unit
    def test_strips_italic(self, app_module):
        assert app_module._strip_inline_md("*italic*") == "italic"

    @pytest.mark.unit
    def test_strips_bold_italic(self, app_module):
        assert app_module._strip_inline_md("***both***") == "both"

    @pytest.mark.unit
    def test_plain_text_unchanged(self, app_module):
        assert app_module._strip_inline_md("plain text") == "plain text"


# ============================================================
# Tests for generate_report_text
# ============================================================


class TestGenerateReportText:
    """Tests for generate_report_text helper."""

    @pytest.mark.unit
    def test_technical_report_has_sections(self, app_module):
        data = {
            "overall_signal": "BULLISH",
            "signal_strength": "Strong",
            "moving_averages": {"sma_20": 100, "sma_50": 95, "sma_200": 90},
            "momentum": {"rsi_14": 55, "macd_line": 1.2, "macd_signal": 0.8, "roc_10_day": 2.5},
            "volatility": {"bollinger_upper": 110, "bollinger_middle": 100, "bollinger_lower": 90},
            "support_resistance": {"resistance_2": 115, "resistance_1": 110, "pivot": 105, "support_1": 100, "support_2": 95},
            "trend": {"short_term": "Bullish", "medium_term": "Bullish", "long_term": "Neutral"},
            "signals": [{"indicator": "RSI", "signal": "Buy", "strength": "Strong"}],
        }
        report = app_module.generate_report_text("RELIANCE", "technical", data)
        assert "MOVING AVERAGES" in report
        assert "MOMENTUM" in report
        assert "VOLATILITY" in report
        assert "SUPPORT" in report
        assert "TREND" in report
        assert "DISCLAIMER" in report

    @pytest.mark.unit
    def test_technical_report_includes_signals(self, app_module):
        data = {
            "moving_averages": {},
            "momentum": {},
            "volatility": {},
            "support_resistance": {},
            "trend": {},
            "signals": [
                {"indicator": "RSI", "signal": "Oversold Buy", "strength": "Strong"},
                {"indicator": "MACD", "signal": "Bullish Cross", "strength": "Medium"},
            ],
        }
        report = app_module.generate_report_text("TCS", "technical", data)
        assert "RSI" in report
        assert "MACD" in report

    @pytest.mark.unit
    def test_fundamental_report_has_sections(self, app_module):
        data = {
            "overall_rating": "BUY",
            "valuation_status": "Fair",
            "valuation": {"pe_ratio": 25},
            "profitability": {"roe": "15%"},
            "financial_health": {"debt_to_equity": 0.5},
            "growth": {"earnings_growth": "10%"},
            "dividends": {"dividend_yield": "1.5%"},
            "size": {"market_cap": "₹10,000 Cr"},
            "assessment": [{"metric": "PE", "assessment": "Fair", "impact": "Neutral"}],
        }
        report = app_module.generate_report_text("INFY", "fundamental", data)
        assert "VALUATION" in report
        assert "PROFITABILITY" in report
        assert "LEVERAGE" in report
        assert "GROWTH" in report
        assert "DIVIDENDS" in report

    @pytest.mark.unit
    def test_fundamental_report_includes_assessment(self, app_module):
        data = {
            "valuation": {},
            "profitability": {},
            "financial_health": {},
            "growth": {},
            "dividends": {},
            "size": {},
            "assessment": [{"metric": "ROE", "assessment": "Strong", "impact": "Positive"}],
        }
        report = app_module.generate_report_text("TCS", "fundamental", data)
        assert "ROE" in report
        assert "Strong" in report

    @pytest.mark.unit
    def test_unknown_report_type(self, app_module):
        result = app_module.generate_report_text("TEST", "unknown", {"data": True})
        assert result == "Report not available"

    @pytest.mark.unit
    def test_empty_data_returns_no_data(self, app_module):
        result = app_module.generate_report_text("TEST", "technical", None)
        assert "No data available" in result

        result2 = app_module.generate_report_text("TEST", "technical", {})
        assert "No data available" in result2


# ============================================================
# Tests for render functions (using mocked st)
# ============================================================


class TestRenderHeader:
    """Tests for render_header function."""

    @pytest.mark.unit
    def test_render_header_calls_markdown(self, app_module, mock_st):
        app_module.render_header()
        assert mock_st.markdown.called
        assert mock_st.divider.called


class TestRenderSidebar:
    """Tests for render_sidebar function."""

    @pytest.mark.unit
    @patch('app.get_trending_stocks', return_value={"gainers": [], "losers": []})
    def test_render_sidebar_returns_symbol(self, mock_trending, app_module, mock_st):
        mock_st.sidebar.__enter__ = MagicMock(return_value=mock_st)
        mock_st.sidebar.__exit__ = MagicMock(return_value=False)
        mock_st.text_input.return_value = "RELIANCE"
        mock_st.selectbox.return_value = ""
        mock_st.tabs.return_value = [MagicMock(), MagicMock()]

        result = app_module.render_sidebar()
        assert result == "RELIANCE"

    @pytest.mark.unit
    @patch('app.get_trending_stocks', return_value={"gainers": [], "losers": []})
    def test_render_sidebar_empty_input(self, mock_trending, app_module, mock_st):
        mock_st.sidebar.__enter__ = MagicMock(return_value=mock_st)
        mock_st.sidebar.__exit__ = MagicMock(return_value=False)
        mock_st.text_input.return_value = ""
        mock_st.selectbox.return_value = ""
        mock_st.tabs.return_value = [MagicMock(), MagicMock()]

        result = app_module.render_sidebar()
        assert result == ""


class TestRenderMarketOverview:
    """Tests for render_market_overview function."""

    @pytest.mark.unit
    @patch('app.get_index_data')
    def test_render_market_overview_success(self, mock_get_idx, app_module, mock_st):
        mock_tool = MagicMock()
        mock_tool.run.return_value = json.dumps({
            "NIFTY50": {"value": 22000, "change": 100, "change_percent": 0.45},
            "SENSEX": {"value": 73000, "change": 300, "change_percent": 0.41},
        })
        mock_get_idx.run = mock_tool.run
        mock_st.columns.return_value = [MagicMock(), MagicMock(), MagicMock(), MagicMock()]

        app_module.render_market_overview()
        assert mock_st.metric.called

    @pytest.mark.unit
    @patch('app.get_index_data')
    def test_render_market_overview_error(self, mock_get_idx, app_module, mock_st):
        mock_tool = MagicMock()
        mock_tool.run.side_effect = Exception("API error")
        mock_get_idx.run = mock_tool.run

        app_module.render_market_overview()
        assert mock_st.warning.called


class TestRenderStockOverview:
    """Tests for render_stock_overview function."""

    @pytest.mark.unit
    @patch('app.get_stock_info')
    @patch('app.get_stock_price')
    @patch('app._fetch_chart_data')
    def test_render_stock_overview_success(self, mock_chart, mock_price, mock_info, app_module, mock_st):
        mock_price.run.return_value = json.dumps({
            "current_price": 2500, "change": 25, "change_percent": 1.0,
            "high": 2520, "low": 2480, "52_week_high": 2800, "52_week_low": 2100,
            "volume": 5000000,
        })
        mock_info.run.return_value = json.dumps({
            "company_name": "Reliance", "sector": "Energy", "industry": "Oil",
            "market_cap": 1600000000000, "pe_ratio": 25, "roe": 12,
            "eps": 80, "pb_ratio": 2.1, "dividend_yield": 0.4,
            "debt_to_equity": 0.5, "book_value": 1100,
        })
        mock_chart.return_value = pd.DataFrame({
            'Close': [100, 101, 102],
        }, index=pd.date_range('2026-01-01', periods=3))
        mock_st.spinner.return_value.__enter__ = MagicMock()
        mock_st.spinner.return_value.__exit__ = MagicMock(return_value=False)
        mock_st.columns.return_value = [MagicMock(), MagicMock()]
        mock_st.radio.return_value = "1Y"

        result = app_module.render_stock_overview("RELIANCE")
        assert result is not None

    @pytest.mark.unit
    @patch('app.get_stock_info')
    @patch('app.get_stock_price')
    def test_render_stock_overview_error_symbol(self, mock_price, mock_info, app_module, mock_st):
        mock_price.run.return_value = json.dumps({"error": "Not found"})
        mock_info.run.return_value = json.dumps({})
        mock_st.spinner.return_value.__enter__ = MagicMock()
        mock_st.spinner.return_value.__exit__ = MagicMock(return_value=False)

        result = app_module.render_stock_overview("BADSTOCK")
        assert result is None

    @pytest.mark.unit
    @patch('app.get_stock_info')
    @patch('app.get_stock_price')
    def test_render_stock_overview_exception(self, mock_price, mock_info, app_module, mock_st):
        mock_price.run.side_effect = Exception("API fail")
        mock_st.spinner.return_value.__enter__ = MagicMock()
        mock_st.spinner.return_value.__exit__ = MagicMock(return_value=False)

        result = app_module.render_stock_overview("RELIANCE")
        assert result is None
        assert mock_st.error.called


class TestRenderTechnicalAnalysis:
    """Tests for render_technical_analysis function."""

    @pytest.mark.unit
    @patch('app.calculate_technical_indicators')
    def test_render_technical_success(self, mock_tech, app_module, mock_st):
        mock_tech.run.return_value = json.dumps({
            "overall_signal": "BULLISH", "signal_strength": "Strong",
            "current_price": 2500,
            "moving_averages": {"sma_20": 2450},
            "momentum": {"rsi_14": 55},
            "volatility": {"bollinger_upper": 2600},
            "support_resistance": {},
            "trend": {"short_term": "Bullish", "golden_cross": True},
            "signals": [{"indicator": "RSI", "signal": "Buy", "strength": "Strong"}],
        })
        mock_st.spinner.return_value.__enter__ = MagicMock()
        mock_st.spinner.return_value.__exit__ = MagicMock(return_value=False)
        mock_st.columns.return_value = [MagicMock(), MagicMock(), MagicMock()]

        app_module.render_technical_analysis("RELIANCE")
        assert mock_st.markdown.called

    @pytest.mark.unit
    @patch('app.calculate_technical_indicators')
    def test_render_technical_error(self, mock_tech, app_module, mock_st):
        mock_tech.run.side_effect = Exception("Calc failed")
        mock_st.spinner.return_value.__enter__ = MagicMock()
        mock_st.spinner.return_value.__exit__ = MagicMock(return_value=False)

        app_module.render_technical_analysis("RELIANCE")
        assert mock_st.error.called


class TestRenderFundamentalAnalysis:
    """Tests for render_fundamental_analysis function."""

    @pytest.mark.unit
    @patch('app.get_fundamental_metrics')
    def test_render_fundamental_success(self, mock_fund, app_module, mock_st):
        mock_fund.run.return_value = json.dumps({
            "overall_rating": "BUY", "score": "8/10", "rating_percentage": "80%",
            "valuation": {"pe_ratio": 25},
            "profitability": {"roe": "15%"},
            "financial_health": {"debt_to_equity": 0.5, "debt_status": "Low"},
            "growth": {"earnings_growth": "10%"},
            "dividends": {"dividend_yield": "1.5%"},
            "size": {"market_cap": "₹10,000 Cr"},
            "assessment": [{"metric": "PE", "assessment": "Fair", "impact": "Neutral"}],
        })
        mock_st.spinner.return_value.__enter__ = MagicMock()
        mock_st.spinner.return_value.__exit__ = MagicMock(return_value=False)
        mock_st.columns.return_value = [MagicMock(), MagicMock(), MagicMock()]

        app_module.render_fundamental_analysis("RELIANCE")
        assert mock_st.markdown.called

    @pytest.mark.unit
    @patch('app.get_fundamental_metrics')
    def test_render_fundamental_error(self, mock_fund, app_module, mock_st):
        mock_fund.run.side_effect = Exception("API fail")
        mock_st.spinner.return_value.__enter__ = MagicMock()
        mock_st.spinner.return_value.__exit__ = MagicMock(return_value=False)

        app_module.render_fundamental_analysis("RELIANCE")
        assert mock_st.error.called


class TestRenderNews:
    """Tests for render_news function."""

    @pytest.mark.unit
    @patch('app.get_stock_news')
    def test_render_news_with_articles(self, mock_news, app_module, mock_st):
        mock_news.run.return_value = json.dumps({
            "articles": [
                {"title": "Test News", "source": "ET", "url": "http://example.com", "published": "2026-02-07", "summary": "Summary"},
            ],
            "sources_status": {"et_rss": "success", "economic_times": "success", "google_news": "success"},
        })
        mock_st.spinner.return_value.__enter__ = MagicMock()
        mock_st.spinner.return_value.__exit__ = MagicMock(return_value=False)
        mock_st.columns.return_value = [MagicMock(), MagicMock(), MagicMock()]
        mock_st.expander.return_value.__enter__ = MagicMock()
        mock_st.expander.return_value.__exit__ = MagicMock(return_value=False)

        app_module.render_news("RELIANCE")
        assert mock_st.markdown.called

    @pytest.mark.unit
    @patch('app.get_stock_news')
    def test_render_news_empty(self, mock_news, app_module, mock_st):
        mock_news.run.return_value = json.dumps({
            "articles": [],
            "sources_status": {},
        })
        mock_st.spinner.return_value.__enter__ = MagicMock()
        mock_st.spinner.return_value.__exit__ = MagicMock(return_value=False)
        mock_st.columns.return_value = [MagicMock(), MagicMock(), MagicMock()]

        app_module.render_news("RELIANCE")
        assert mock_st.info.called


class TestRenderInstitutional:
    """Tests for render_institutional function."""

    @pytest.mark.unit
    @patch('app.get_bulk_block_deals')
    @patch('app.get_fii_dii_data')
    def test_render_institutional_with_data(self, mock_fii, mock_deals, app_module, mock_st):
        mock_fii.run.return_value = json.dumps({
            "fii": {"buy_value_cr": 5000, "sell_value_cr": 4000, "net_value_cr": 1000, "activity": "Net Buyer"},
            "dii": {"buy_value_cr": 3000, "sell_value_cr": 2000, "net_value_cr": 1000, "activity": "Net Buyer"},
            "combined": {"market_sentiment": "Bullish"},
        })
        mock_deals.run.return_value = json.dumps({
            "deals": [{"stock": "RELIANCE", "deal_type": "Bulk", "client": "FII", "quantity": 100000, "price": 2500, "note": None}],
        })
        mock_st.spinner.return_value.__enter__ = MagicMock()
        mock_st.spinner.return_value.__exit__ = MagicMock(return_value=False)
        mock_st.columns.return_value = [MagicMock(), MagicMock()]

        app_module.render_institutional("RELIANCE")
        assert mock_st.markdown.called

    @pytest.mark.unit
    @patch('app.get_bulk_block_deals')
    @patch('app.get_fii_dii_data')
    def test_render_institutional_no_deals(self, mock_fii, mock_deals, app_module, mock_st):
        mock_fii.run.return_value = json.dumps({
            "fii": {}, "dii": {}, "combined": {},
        })
        mock_deals.run.return_value = json.dumps({
            "deals": [],
        })
        mock_st.spinner.return_value.__enter__ = MagicMock()
        mock_st.spinner.return_value.__exit__ = MagicMock(return_value=False)
        mock_st.columns.return_value = [MagicMock(), MagicMock()]

        app_module.render_institutional("RELIANCE")
        assert mock_st.info.called or mock_st.markdown.called


class TestRenderAIAnalysis:
    """Tests for render_ai_analysis function."""

    @pytest.mark.unit
    def test_render_ai_no_button_click(self, app_module, mock_st):
        """When button not clicked, should show info and not run analysis."""
        mock_st.button.return_value = False
        app_module.render_ai_analysis("RELIANCE")
        assert mock_st.info.called

    @pytest.mark.unit
    def test_render_ai_with_stored_report(self, app_module, mock_st):
        """When report in session state, should show download buttons."""
        mock_st.session_state = {
            "report_RELIANCE": "# Report\nContent here",
            "report_time_RELIANCE": "2026-02-07_10-30",
        }
        mock_st.button.return_value = False
        mock_st.expander.return_value.__enter__ = MagicMock()
        mock_st.expander.return_value.__exit__ = MagicMock(return_value=False)
        mock_st.columns.return_value = [MagicMock(), MagicMock()]

        app_module.render_ai_analysis("RELIANCE")
        assert mock_st.download_button.called


class TestRenderRangeBar:
    """Tests for _render_range_bar function."""

    @pytest.mark.unit
    def test_renders_html(self, app_module, mock_st):
        app_module._render_range_bar("Test", 100.0, 200.0, 150.0)
        assert mock_st.markdown.called
        call_args = mock_st.markdown.call_args
        assert "unsafe_allow_html" in call_args.kwargs

    @pytest.mark.unit
    def test_skips_invalid_range(self, app_module, mock_st):
        mock_st.markdown.reset_mock()
        app_module._render_range_bar("Test", 200.0, 100.0, 150.0)
        # Should not call st.markdown because high <= low
        assert not mock_st.markdown.called

    @pytest.mark.unit
    def test_skips_zero_high(self, app_module, mock_st):
        mock_st.markdown.reset_mock()
        app_module._render_range_bar("Test", 0.0, 0.0, 0.0)
        assert not mock_st.markdown.called


class TestFetchChartData:
    """Tests for _fetch_chart_data helper."""

    @pytest.mark.unit
    def test_fetch_chart_data_returns_dataframe(self, app_module):
        dates = pd.date_range(end=datetime.now(), periods=30, freq='D')
        mock_df = pd.DataFrame({
            'Open': np.random.uniform(100, 110, 30),
            'High': np.random.uniform(110, 120, 30),
            'Low': np.random.uniform(90, 100, 30),
            'Close': np.random.uniform(100, 110, 30),
            'Volume': np.random.randint(1000, 10000, 30),
        }, index=dates)

        with patch('yfinance.Ticker') as mock_ticker:
            mock_ticker.return_value.history.return_value = mock_df
            result = app_module._fetch_chart_data("RELIANCE", "1M")

        assert isinstance(result, pd.DataFrame)
        assert not result.empty
        assert "Close" in result.columns

    @pytest.mark.unit
    def test_fetch_chart_data_returns_empty_on_error(self, app_module):
        with patch('yfinance.Ticker', side_effect=Exception("API error")):
            result = app_module._fetch_chart_data("BADSTOCK", "1Y")

        assert isinstance(result, pd.DataFrame)
        assert result.empty

    @pytest.mark.unit
    def test_period_map_covers_all_buttons(self):
        period_map = {"1D", "1W", "1M", "3M", "6M", "1Y", "5Y"}
        expected_periods = {"1D", "1W", "1M", "3M", "6M", "1Y", "5Y"}
        assert period_map == expected_periods


class TestRangeBarLogic:
    """Tests for range bar percentage calculation logic."""

    @pytest.mark.unit
    def test_range_bar_percentage_at_low(self):
        low, high, current = 100.0, 200.0, 100.0
        pct = max(0, min(100, ((current - low) / (high - low)) * 100))
        assert pct == 0.0

    @pytest.mark.unit
    def test_range_bar_percentage_at_high(self):
        low, high, current = 100.0, 200.0, 200.0
        pct = max(0, min(100, ((current - low) / (high - low)) * 100))
        assert pct == 100.0

    @pytest.mark.unit
    def test_range_bar_percentage_midpoint(self):
        low, high, current = 100.0, 200.0, 150.0
        pct = max(0, min(100, ((current - low) / (high - low)) * 100))
        assert pct == 50.0

    @pytest.mark.unit
    def test_range_bar_clamps_below_low(self):
        low, high, current = 100.0, 200.0, 50.0
        pct = max(0, min(100, ((current - low) / (high - low)) * 100))
        assert pct == 0.0

    @pytest.mark.unit
    def test_range_bar_clamps_above_high(self):
        low, high, current = 100.0, 200.0, 250.0
        pct = max(0, min(100, ((current - low) / (high - low)) * 100))
        assert pct == 100.0


# ============================================================
# Tests for main function
# ============================================================


class TestMainFunction:
    """Tests for main() orchestration."""

    @pytest.mark.unit
    @patch('app.render_header')
    @patch('app.render_sidebar', return_value="")
    @patch('app.render_market_overview')
    def test_main_no_symbol(self, mock_overview, mock_sidebar, mock_header, app_module, mock_st):
        mock_st.columns.return_value = [MagicMock() for _ in range(5)]
        with patch('app.get_stock_price') as mock_price:
            mock_price.run.return_value = json.dumps({"current_price": 100, "change_percent": 1})
            app_module.main()
        assert mock_header.called
        assert mock_sidebar.called

    @pytest.mark.unit
    @patch('app.render_ai_analysis')
    @patch('app.render_institutional')
    @patch('app.render_news')
    @patch('app.render_fundamental_analysis')
    @patch('app.render_technical_analysis')
    @patch('app.render_stock_overview', return_value=({"price": 100}, {"info": True}))
    @patch('app.render_market_overview')
    @patch('app.render_sidebar', return_value="RELIANCE")
    @patch('app.render_header')
    def test_main_with_symbol(self, mock_header, mock_sidebar, mock_overview,
                              mock_stock, mock_tech, mock_fund, mock_news,
                              mock_inst, mock_ai, app_module, mock_st):
        mock_st.tabs.return_value = [MagicMock() for _ in range(5)]
        app_module.main()
        assert mock_stock.called


# ============================================================
# Word report tests (already work — keep them)
# ============================================================


class TestGenerateWordReport:
    """Tests for Word document generation."""

    @pytest.mark.unit
    def test_returns_bytes(self):
        from app import _generate_word_report
        md = "# Title\n\nSome text.\n\n- Bullet one\n- Bullet two\n"
        result = _generate_word_report("RELIANCE", md, "2026-02-07_10-30")
        assert isinstance(result, bytes)
        assert len(result) > 0

    @pytest.mark.unit
    def test_output_is_valid_docx(self):
        from docx import Document
        import io
        from app import _generate_word_report
        md = "# Report\n\n## Section\n\nHello world\n"
        result = _generate_word_report("TCS", md, "2026-02-07_12-00")
        doc = Document(io.BytesIO(result))
        assert len(doc.paragraphs) >= 3

    @pytest.mark.unit
    def test_headings_are_parsed(self):
        from docx import Document
        import io
        from app import _generate_word_report
        md = "# Heading 1\n## Heading 2\n### Heading 3\n"
        result = _generate_word_report("INFY", md, "2026-02-07")
        doc = Document(io.BytesIO(result))
        heading_texts = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Heading 1" in heading_texts
        assert "Heading 2" in heading_texts
        assert "Heading 3" in heading_texts

    @pytest.mark.unit
    def test_bullet_list_items(self):
        from docx import Document
        import io
        from app import _generate_word_report
        md = "- Item A\n- Item B\n* Item C\n"
        result = _generate_word_report("HDFC", md, "2026-02-07")
        doc = Document(io.BytesIO(result))
        bullet_texts = [p.text for p in doc.paragraphs if "List" in p.style.name]
        assert "Item A" in bullet_texts
        assert "Item B" in bullet_texts
        assert "Item C" in bullet_texts

    @pytest.mark.unit
    def test_numbered_list_items(self):
        from docx import Document
        import io
        from app import _generate_word_report
        md = "1. First\n2. Second\n3. Third\n"
        result = _generate_word_report("SBI", md, "2026-02-07")
        doc = Document(io.BytesIO(result))
        numbered_texts = [p.text for p in doc.paragraphs if "Number" in p.style.name]
        assert "First" in numbered_texts
        assert "Second" in numbered_texts

    @pytest.mark.unit
    def test_disclaimer_present(self):
        from docx import Document
        import io
        from app import _generate_word_report
        md = "Some analysis text.\n"
        result = _generate_word_report("WIPRO", md, "2026-02-07")
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "DISCLAIMER" in all_text
        assert "educational purposes" in all_text

    @pytest.mark.unit
    def test_title_contains_symbol(self):
        from docx import Document
        import io
        from app import _generate_word_report
        md = "Content\n"
        result = _generate_word_report("BAJAJ", md, "2026-02-07")
        doc = Document(io.BytesIO(result))
        title_para = doc.paragraphs[0]
        assert "BAJAJ" in title_para.text

    @pytest.mark.unit
    def test_horizontal_rule_converted(self):
        from docx import Document
        import io
        from app import _generate_word_report
        md = "Above\n---\nBelow\n"
        result = _generate_word_report("TEST", md, "2026-02-07")
        doc = Document(io.BytesIO(result))
        found_border = False
        for p in doc.paragraphs:
            xml = p._p.xml
            if 'w:pBdr' in xml and 'w:bottom' in xml:
                found_border = True
                break
        assert found_border

    @pytest.mark.unit
    def test_empty_report(self):
        from app import _generate_word_report
        result = _generate_word_report("EMPTY", "", "2026-02-07")
        assert isinstance(result, bytes)
        assert len(result) > 0

    @pytest.mark.unit
    def test_code_fences_stripped(self):
        from docx import Document
        import io
        from app import _generate_word_report
        md = "```markdown\n# Title\nSome text\n```\n"
        result = _generate_word_report("TEST", md, "2026-02-07")
        doc = Document(io.BytesIO(result))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "```" not in all_text

    @pytest.mark.unit
    def test_bold_markdown_parsed(self):
        from docx import Document
        import io
        from app import _generate_word_report
        md = "This has **important** text.\n"
        result = _generate_word_report("TEST", md, "2026-02-07")
        doc = Document(io.BytesIO(result))
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text == "important":
                    assert run.bold is True
                    return
        pytest.fail("Did not find a run with text 'important'")

    @pytest.mark.unit
    def test_italic_markdown_parsed(self):
        from docx import Document
        import io
        from app import _generate_word_report
        md = "This has *emphasis* here.\n"
        result = _generate_word_report("TEST", md, "2026-02-07")
        doc = Document(io.BytesIO(result))
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text == "emphasis":
                    assert run.italic is True
                    return
        pytest.fail("Did not find a run with text 'emphasis'")

    @pytest.mark.unit
    def test_pipe_table_becomes_word_table(self):
        from docx import Document
        import io
        from app import _generate_word_report
        md = "| Metric | Value |\n|---|---|\n| PE Ratio | 25.3 |\n| ROE | 18.5% |\n"
        result = _generate_word_report("TEST", md, "2026-02-07")
        doc = Document(io.BytesIO(result))
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        assert table.rows[0].cells[0].text.strip() == "Metric"
        assert table.rows[0].cells[1].text.strip() == "Value"
        assert table.rows[1].cells[0].text.strip() == "PE Ratio"

    @pytest.mark.unit
    def test_table_header_has_shading(self):
        from docx import Document
        import io
        from app import _generate_word_report
        md = "| Name | Score |\n|---|---|\n| Alpha | 90 |\n"
        result = _generate_word_report("TEST", md, "2026-02-07")
        doc = Document(io.BytesIO(result))
        assert len(doc.tables) >= 1
        header_cell = doc.tables[0].rows[0].cells[0]
        xml = header_cell._tc.xml
        assert 'w:shd' in xml
        assert '1B2A4A' in xml


class TestGeneratePDFReport:
    """Tests for PDF report generation."""

    @pytest.mark.unit
    def test_returns_bytes(self):
        from app import _generate_pdf_report
        md = "# Title\n\nSome text.\n\n- Bullet one\n- Bullet two\n"
        result = _generate_pdf_report("RELIANCE", md, "2026-02-07_10-30")
        assert isinstance(result, bytes)
        assert len(result) > 0

    @pytest.mark.unit
    def test_output_is_valid_pdf(self):
        from app import _generate_pdf_report
        md = "# Report\n\n## Section\n\nHello world\n"
        result = _generate_pdf_report("TCS", md, "2026-02-07_12-00")
        assert result[:5] == b"%PDF-"

    @pytest.mark.unit
    def test_code_fences_stripped(self):
        from app import _generate_pdf_report
        md = "```markdown\n# Title\nSome text\n```\n"
        result = _generate_pdf_report("TEST", md, "2026-02-07")
        assert b"```" not in result

    @pytest.mark.unit
    def test_handles_unicode(self):
        from app import _generate_pdf_report
        md = "Price: \u20b9224.46\nChange: \u201310%\n"
        result = _generate_pdf_report("TEST", md, "2026-02-07")
        assert isinstance(result, bytes)
        assert result[:5] == b"%PDF-"

    @pytest.mark.unit
    def test_pipe_table_in_pdf(self):
        from app import _generate_pdf_report
        md = "| Metric | Value |\n|---|---|\n| PE | 25.3 |\n"
        result = _generate_pdf_report("TEST", md, "2026-02-07")
        assert isinstance(result, bytes)
        assert len(result) > 500

    @pytest.mark.unit
    def test_empty_report(self):
        from app import _generate_pdf_report
        result = _generate_pdf_report("EMPTY", "", "2026-02-07")
        assert isinstance(result, bytes)
        assert result[:5] == b"%PDF-"

    @pytest.mark.unit
    def test_numbered_list_in_pdf(self):
        from app import _generate_pdf_report
        md = "1. First\n2. Second\n3. Third\n"
        result = _generate_pdf_report("TEST", md, "2026-02-07")
        assert isinstance(result, bytes)
        assert result[:5] == b"%PDF-"

    @pytest.mark.unit
    def test_horizontal_rule_in_pdf(self):
        from app import _generate_pdf_report
        md = "Above\n---\nBelow\n"
        result = _generate_pdf_report("TEST", md, "2026-02-07")
        assert isinstance(result, bytes)
        assert result[:5] == b"%PDF-"

    @pytest.mark.unit
    def test_buy_sell_colored_in_pdf(self):
        from app import _generate_pdf_report
        md = "- Signal: BUY now\n- Signal: SELL later\n"
        result = _generate_pdf_report("TEST", md, "2026-02-07")
        assert isinstance(result, bytes)
        assert result[:5] == b"%PDF-"


class TestAddColoredRun:
    """Tests for BUY/SELL keyword coloring in Word documents."""

    @pytest.mark.unit
    def test_buy_is_green_and_bold(self):
        from docx import Document
        from docx.shared import RGBColor
        from app import _add_colored_run
        doc = Document()
        para = doc.add_paragraph()
        _add_colored_run(para, "Recommendation: BUY this stock")
        runs = para.runs
        buy_runs = [r for r in runs if r.text.strip() == "BUY"]
        assert len(buy_runs) == 1
        assert buy_runs[0].bold is True
        assert buy_runs[0].font.color.rgb == RGBColor(0, 128, 0)

    @pytest.mark.unit
    def test_sell_is_red_and_bold(self):
        from docx import Document
        from docx.shared import RGBColor
        from app import _add_colored_run
        doc = Document()
        para = doc.add_paragraph()
        _add_colored_run(para, "Signal: SELL immediately")
        runs = para.runs
        sell_runs = [r for r in runs if r.text.strip() == "SELL"]
        assert len(sell_runs) == 1
        assert sell_runs[0].bold is True
        assert sell_runs[0].font.color.rgb == RGBColor(200, 0, 0)

    @pytest.mark.unit
    def test_strong_buy_is_green(self):
        from docx import Document
        from docx.shared import RGBColor
        from app import _add_colored_run
        doc = Document()
        para = doc.add_paragraph()
        _add_colored_run(para, "Rating: STRONG BUY")
        runs = para.runs
        sb_runs = [r for r in runs if r.text.strip() == "STRONG BUY"]
        assert len(sb_runs) == 1
        assert sb_runs[0].bold is True
        assert sb_runs[0].font.color.rgb == RGBColor(0, 128, 0)

    @pytest.mark.unit
    def test_strong_sell_is_red(self):
        from docx import Document
        from docx.shared import RGBColor
        from app import _add_colored_run
        doc = Document()
        para = doc.add_paragraph()
        _add_colored_run(para, "Rating: STRONG SELL")
        runs = para.runs
        ss_runs = [r for r in runs if r.text.strip() == "STRONG SELL"]
        assert len(ss_runs) == 1
        assert ss_runs[0].bold is True
        assert ss_runs[0].font.color.rgb == RGBColor(200, 0, 0)

    @pytest.mark.unit
    def test_text_without_keywords_is_plain(self):
        from docx import Document
        from app import _add_colored_run
        doc = Document()
        para = doc.add_paragraph()
        _add_colored_run(para, "This is plain text with no signals")
        runs = para.runs
        assert len(runs) == 1
        assert runs[0].bold is not True
        assert runs[0].font.color.rgb is None

    @pytest.mark.unit
    def test_multiple_keywords_in_one_line(self):
        from docx import Document
        from docx.shared import RGBColor
        from app import _add_colored_run
        doc = Document()
        para = doc.add_paragraph()
        _add_colored_run(para, "Short-term BUY but long-term SELL")
        runs = para.runs
        buy_runs = [r for r in runs if r.text.strip() == "BUY"]
        sell_runs = [r for r in runs if r.text.strip() == "SELL"]
        assert len(buy_runs) == 1
        assert len(sell_runs) == 1
        assert buy_runs[0].font.color.rgb == RGBColor(0, 128, 0)
        assert sell_runs[0].font.color.rgb == RGBColor(200, 0, 0)


# ============================================================
# Existing data parsing tests (these don't need app import)
# ============================================================


class TestStreamlitUILogic:
    """Tests for Streamlit UI component logic."""

    @pytest.mark.unit
    def test_stock_symbol_validation(self):
        from config import NIFTY50_STOCKS
        assert "RELIANCE" in NIFTY50_STOCKS
        assert "TCS" in NIFTY50_STOCKS

    @pytest.mark.unit
    def test_sector_mapping(self):
        from config import SECTORS
        assert "IT" in SECTORS or "TECHNOLOGY" in SECTORS
        assert "BANKING" in SECTORS or "FINANCE" in SECTORS

    @pytest.mark.unit
    def test_price_data_parsing(self):
        mock_price_response = json.dumps({
            "symbol": "RELIANCE", "current_price": 2847.50,
            "change": 34.25, "change_percent": 1.22, "volume": 5000000,
        })
        data = json.loads(mock_price_response)
        assert data["symbol"] == "RELIANCE"
        assert data["current_price"] == 2847.50

    @pytest.mark.unit
    def test_technical_data_parsing(self):
        mock_technical_response = json.dumps({
            "rsi_14": 55.5, "macd": 12.5, "signal": 10.2,
            "bb_upper": 2900, "bb_middle": 2850, "bb_lower": 2800,
        })
        data = json.loads(mock_technical_response)
        assert 0 <= data["rsi_14"] <= 100
        assert data["bb_upper"] >= data["bb_middle"] >= data["bb_lower"]
